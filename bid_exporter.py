#!/usr/bin/env python
# -*- coding: utf-8 -*-
"""
标书导出模块
支持导出为Word、PDF等格式
"""
import io
import sys
import re
from datetime import datetime
from typing import Dict, Optional

try:
    from docx import Document
    from docx.shared import Pt, RGBColor, Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False
    print("警告：python-docx未安装，Word导出功能不可用")

class BidExporter:
    """标书导出器"""

    def __init__(self):
        """初始化"""
        if not DOCX_AVAILABLE:
            raise ImportError("请先安装python-docx: pip install python-docx")

    def markdown_to_word(self, markdown_content: str, output_path: str, metadata: Dict = None) -> bool:
        """将Markdown标书转换为Word文档"""
        if not DOCX_AVAILABLE:
            print("错误：python-docx未安装")
            return False

        try:
            doc = Document()

            # 设置默认字体
            doc.styles['Normal'].font.name = '宋体'
            doc.styles['Normal']._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
            doc.styles['Normal'].font.size = Pt(12)

            # 添加元数据
            if metadata:
                self._add_metadata(doc, metadata)

            # 解析Markdown
            lines = markdown_content.split('\n')
            i = 0

            while i < len(lines):
                line = lines[i].rstrip()

                # 标题
                if line.startswith('# '):
                    self._add_heading(doc, line[2:], level=1)
                elif line.startswith('## '):
                    self._add_heading(doc, line[3:], level=2)
                elif line.startswith('### '):
                    self._add_heading(doc, line[4:], level=3)
                elif line.startswith('#### '):
                    self._add_heading(doc, line[5:], level=4)

                # 列表
                elif line.startswith(('- ', '* ', '• ')):
                    self._add_list_item(doc, line[2:])

                # 编号列表
                elif re.match(r'^\d+\. ', line):
                    self._add_numbered_item(doc, line)

                # 空行
                elif not line.strip():
                    pass  # 跳过空行

                # 普通段落
                elif line.strip():
                    # 合并多行段落
                    paragraph_text = line
                    i += 1
                    while i < len(lines) and lines[i].strip() and not lines[i].startswith(('#', '-', '*', '•')):
                        paragraph_text += ' ' + lines[i].strip()
                        i += 1
                    i -= 1  # 回退一行

                    self._add_paragraph(doc, paragraph_text)

                i += 1

            # 保存文档
            doc.save(output_path)
            print(f"✓ 导出成功: {output_path}")
            return True

        except Exception as e:
            print(f"✗ 导出失败: {e}")
            return False

    def _add_metadata(self, doc: Document, metadata: Dict):
        """添加文档元数据"""
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT

        if metadata.get('project_name'):
            p.add_run(f"项目名称：{metadata['project_name']}\n")

        if metadata.get('company'):
            p.add_run(f"投标单位：{metadata['company']}\n")

        if metadata.get('date'):
            p.add_run(f"日期：{metadata['date']}\n")
        else:
            p.add_run(f"日期：{datetime.now().strftime('%Y年%m月%d日')}\n")

        doc.add_paragraph()  # 空行

    def _add_heading(self, doc: Document, text: str, level: int = 1):
        """添加标题"""
        heading = doc.add_heading(text, level=level)

        # 设置标题样式
        if level == 1:
            heading.runs[0].font.size = Pt(18)
            heading.runs[0].font.bold = True
            heading.runs[0].font.color.rgb = RGBColor(0, 51, 102)
        elif level == 2:
            heading.runs[0].font.size = Pt(16)
            heading.runs[0].font.bold = True
        elif level == 3:
            heading.runs[0].font.size = Pt(14)
            heading.runs[0].font.bold = True

    def _add_paragraph(self, doc: Document, text: str):
        """添加段落"""
        p = doc.add_paragraph(text)
        p.paragraph_format.line_spacing = 1.5
        p.paragraph_format.space_after = Pt(6)

    def _add_list_item(self, doc: Document, text: str):
        """添加无序列表项"""
        p = doc.add_paragraph(text, style='List Bullet')
        p.paragraph_format.line_spacing = 1.5

    def _add_numbered_item(self, doc: Document, text: str):
        """添加有序列表项"""
        p = doc.add_paragraph(text, style='List Number')
        p.paragraph_format.line_spacing = 1.5

    def export_bid_package(self, bid_content: str, metadata: Dict, output_dir: str = ".") -> Dict:
        """导出完整标书包"""
        results = {
            "word": None,
            "markdown": None,
            "success": False
        }

        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        project_name = metadata.get('project_name', '未命名项目')
        safe_name = re.sub(r'[\\/*?:"<>|]', '_', project_name)

        # 导出Markdown
        try:
            md_path = f"{output_dir}/{safe_name}_{timestamp}.md"
            with open(md_path, 'w', encoding='utf-8') as f:
                f.write(f"# {project_name}\n\n")
                if metadata.get('project_description'):
                    f.write(f"## 项目描述\n{metadata['project_description']}\n\n")
                f.write(f"## 标书方案\n\n")
                f.write(bid_content)
            results['markdown'] = md_path
            print(f"✓ Markdown导出: {md_path}")
        except Exception as e:
            print(f"✗ Markdown导出失败: {e}")

        # 导出Word
        if DOCX_AVAILABLE:
            try:
                word_path = f"{output_dir}/{safe_name}_{timestamp}.docx"
                if self.markdown_to_word(bid_content, word_path, metadata):
                    results['word'] = word_path
            except Exception as e:
                print(f"✗ Word导出失败: {e}")

        results['success'] = results['markdown'] is not None or results['word'] is not None
        return results

# 测试
if __name__ == '__main__':
    if DOCX_AVAILABLE:
        exporter = BidExporter()

        sample_bid = """# 智慧城市综合管理平台

## 一、项目理解与需求分析

### 1.1 项目背景
本项目旨在建设集交通管理、环境监测、公共安全等功能于一体的智慧城市综合管理平台。

### 1.2 需求要点
- 系统集成度高
- 实时数据处理
- AI智能分析能力

## 二、技术方案

### 2.1 总体架构
采用微服务架构设计，系统分为：
1. 用户层
2. 应用层
3. 服务层
4. 数据层

### 2.2 技术选型
- 前端：Vue.js 3.0
- 后端：Spring Boot 2.7
- 数据库：MySQL 8.0
- 缓存：Redis 6.0

## 三、实施方案

项目周期12个月，分为四个阶段：
- 需求分析：1个月
- 系统设计：2个月
- 开发实施：6个月
- 测试上线：3个月
"""

        metadata = {
            "project_name": "智慧城市综合管理平台",
            "company": "XX科技有限公司",
            "date": "2026年1月31日",
            "project_description": "建设集交通、环境、安全于一体的智慧城市平台"
        }

        print("=== 标书导出测试 ===\n")
        result = exporter.export_bid_package(sample_bid, metadata)
        print(f"\n导出结果: {'成功' if result['success'] else '失败'}")
    else:
        print("请先安装python-docx: pip install python-docx")
